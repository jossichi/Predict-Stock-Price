import os
from datetime import datetime, timedelta
import pandas as pd
import numpy as np
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import MinMaxScaler
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import LSTM, Dense
from tensorflow.keras.callbacks import EarlyStopping
from sklearn.metrics import mean_squared_error
from docx import Document
from keras.optimizers import Adam
from keras.initializers import glorot_normal
import tensorflow as tf

class LSTMModel:
    
    def __init__(self, file_path, sheet_name, target_column='Close', feature_columns=None, test_size=0.2, current_date=None):
        self.file_path = file_path
        self.sheet_name = sheet_name
        self.target_column = target_column
        self.feature_columns = feature_columns
        self.test_size = test_size
        self.current_date = current_date or datetime.now().strftime('%Y-%m-%d')
        self.df = None
        self.next_day = (datetime.now() + timedelta(days=1)).strftime('%Y-%m-%d')
        self.model = None
        self.merged_df = None
        self.X = None
        self.y = None
        self.retrain_attempts = 3

    def check_exist_data(self):
        if os.path.exists(self.file_path):
            print(f"File '{self.file_path}' exists.")
            excel_sheets = pd.ExcelFile(self.file_path).sheet_names

            for sheet_name in excel_sheets:
                if sheet_name == self.current_date:
                    print(f"\nSheet with current date '{self.current_date}' exists in the file.")
                    if self.load_data(sheet_name):
                        print(f"\nSample of loaded DataFrame for {sheet_name}:")
                        print(self.df.head())

                        consecutive_days = self.get_consecutive_days(excel_sheets)
                        if consecutive_days >4 :
                            print("Using long-term predictions for training.")
                            self.long_terms_prediction()
                        else:
                            print(f"Using short-term predictions for the next day: {self.next_day}")
                            predictions_df=self.short_terms_prediction()
                            print(f"Predicted stock price for the next day: {self.next_day}")                           
                            
                            # self.print_recommendations(predictions, actual_values)
                            self.export_to_word(predictions_df)

                else:
                    print(f"\nSheet with current date '{self.current_date}' does not exist in the file.")

        else:
            print(f"File '{self.file_path}' does not exist.")

    def load_data(self, sheet_name):
        try:
            self.df = pd.read_excel(self.file_path, sheet_name=sheet_name)
            print("dataframe before predtictions:\n")
            print(self.df)
            return True
        except Exception as e:
            print(f"Error loading data for {sheet_name}: {e}")
            return False

    def get_consecutive_days(self, excel_sheets):
        date_index = excel_sheets.index(self.current_date)
        consecutive_days = 1

        for i in range(date_index + 1, len(excel_sheets)):
            next_date = datetime.strptime(excel_sheets[i], '%Y-%m-%d')
            prev_date = datetime.strptime(excel_sheets[i - 1], '%Y-%m-%d')

            if (next_date - prev_date).days == 1:
                consecutive_days += 1
            else:
                break

        return consecutive_days

    def long_terms_prediction(self):
        target_column = self.target_column

        X, y = self.df[self.feature_columns], self.df[target_column]
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=self.test_size, shuffle=False)

        X_train, y_train = self.preprocess_data(X_train, y_train)
        X_test, y_test = self.preprocess_data(X_test, y_test)

        self.train_lstm_model(X_train, y_train)

        predictions = self.model.predict(X_test)

        mse = mean_squared_error(y_test, predictions)
        print(f"Mean Squared Error (MSE) on test data: {mse}")

        self.print_recommendations(predictions, y_test)
        
        return predictions

    def short_terms_prediction(self):
        target_column = self.target_column

        # Create an empty DataFrame to store predictions for each symbol
        predictions_df = pd.DataFrame()

        symbols_with_nan = []
        symbols_with_extreme_changes = []

        # Extract unique symbols from the DataFrame
        symbols = self.df['Symbol'].unique()

        for symbol in symbols:
            previous_percent_change = None
            retrain_attempts = 3  # Set the maximum number of retraining attempts

            # Extract features and target variable for the current symbol
            symbol_data = self.df[self.df['Symbol'] == symbol]
            self.X, self.y = symbol_data[self.feature_columns], symbol_data[target_column]

            while retrain_attempts > 0:
                self.train_lstm_model(self.X, self.y)

                # Take the last row of symbol_data to get the latest available data
                latest_data = symbol_data.tail(1)
                self.X_latest = latest_data[self.feature_columns]  # Store it as a class variable
                next_day_prediction = self.model.predict(np.reshape(self.X_latest.values, (1, self.X_latest.shape[1], 1)))

                # Add the prediction to the DataFrame
                min_prediction = next_day_prediction[0][0]
                percent_change = ((min_prediction - latest_data[target_column].values[0]) / latest_data[target_column].values[0]) * 100

                if np.isnan(next_day_prediction[0][0]):
                    # Check if the prediction contains NaN before rounding
                    print(f"Symbol {symbol} contains NaN in the prediction. Retraining...")
                    symbols_with_nan.append(symbol)
                    retrain_attempts -= 1
                    continue

                percent_change = round(percent_change)

                if -100 < percent_change < 100:
                    predictions_df = pd.concat([predictions_df, pd.DataFrame({
                        'Symbol': [symbol],
                        'Prediction': [min_prediction],
                        'Actual Value': [latest_data[self.target_column].values[0]],
                        'Predicted_Percentage_Change': [percent_change]
                    })])
                    break

                elif percent_change > 100 or percent_change < -100:
                    print(f"Symbol {symbol} has percent change is outta range (-100;100) Retraining...")
                    symbols_with_extreme_changes.append(symbol)
                    retrain_attempts -= 1

                    # Retrain the model for the specific symbol
                    self.X, self.y = symbol_data[self.feature_columns], symbol_data[target_column]
                    previous_percent_change = percent_change
                    continue

            if retrain_attempts == 0:
                print(f"Exceeded maximum retraining attempts for symbol {symbol}. Unable to obtain desired percent change.")

        return predictions_df

    def export_to_word(self, predictions_df):
        document = Document()
        document.add_heading('Stock Predictions Report', 0)
        
        print(predictions_df.columns)
        
        # Add content to the Word document
        document.add_paragraph(f"Predicted stock prices for all symbols on the next day ({self.next_day}):")
        table = document.add_table(rows=len(predictions_df) + 1, cols=6)  # Assuming six columns

        # Write header row
        table.cell(0, 0).text = 'Symbol'
        table.cell(0, 1).text = 'Prediction'
        table.cell(0, 2).text = 'Actual Value'  

        table.cell(0, 3).text = 'Percentage Change'
        table.cell(0, 4).text = 'Absolute Change'
        table.cell(0, 5).text = 'Recommendation'

        # Write data rows
        for i, (symbol, prediction, actual_value, percentage_change) in enumerate(zip(predictions_df['Symbol'], predictions_df['Prediction'], predictions_df['Actual Value'], predictions_df['Predicted_Percentage_Change'])):
            table.cell(i + 1, 0).text = str(symbol)
            table.cell(i + 1, 1).text = str(prediction)
            table.cell(i + 1, 2).text = str(actual_value)
            table.cell(i + 1, 3).text = str(percentage_change)
            # Calculate absolute change
            absolute_change = prediction - actual_value
            table.cell(i + 1, 4).text = f"{absolute_change:.2f}"

            # Include the recommendation based on conditions
            threshold = 0.01  # Adjust threshold based on your preference
            if prediction is not None and actual_value is not None:
                if prediction > actual_value + threshold:
                    recommendation = 'Buy'
                elif prediction < actual_value - threshold:
                    recommendation = 'Sell'
                else:
                    recommendation = 'Hold'
                table.cell(i + 1, 5).text = recommendation

        # Save the document
        report_path = 'report.docx'
        document.save(report_path)
        print(f"Report saved to '{report_path}'.")

    def train_lstm_model(self, X_train, y_train, epochs=50, batch_size=32, learning_rate=0.001):
        if X_train.shape[0] < 3:
            print("Warning: Insufficient data for validation split. Using the entire dataset for training.")
            validation_split = 0
        else:
            validation_split = 0.2

        self.model = Sequential()
        self.model.add(LSTM(50, activation='relu', input_shape=(X_train.shape[1], 1)))
        self.model.add(Dense(1, kernel_initializer=glorot_normal(seed=None)))  # Random initialization

        optimizer = Adam(learning_rate=learning_rate)
        self.model.compile(optimizer=optimizer, loss='mse', metrics=[self.mean_absolute_percentage_error])

        early_stop = EarlyStopping(monitor='val_loss', patience=5, restore_best_weights=True)

        X_train = np.reshape(X_train, (X_train.shape[0], X_train.shape[1], 1))

        self.model.fit(X_train, y_train, epochs=epochs, batch_size=batch_size, validation_split=validation_split, callbacks=[early_stop])
    
    def mean_absolute_percentage_error(self, y_true, y_pred):
        y_true = tf.squeeze(y_true)
        y_pred = tf.squeeze(y_pred)
        return tf.reduce_mean(tf.abs((y_true - y_pred) / tf.maximum(tf.abs(y_true), tf.keras.backend.epsilon()))) * 100
    
    def preprocess_data(self, X, y):
        scaler = MinMaxScaler()
        X_scaled = scaler.fit_transform(X)
        y_scaled = scaler.fit_transform(np.array(y).reshape(-1, 1))

        X_reshaped = np.reshape(X_scaled, (X_scaled.shape[0], X_scaled.shape[1], 1))

        return X_reshaped, y_scaled
    
current_date = datetime.now().strftime('%Y-%m-%d')
file_path = "stock_data.xlsx"
sheet_name = current_date
feature_columns = ['Open', 'High', 'Low', 'Volume', 'RSI', 'MACD_diff', 'BB_upper', 'BB_lower', 'Returns_daily', 'Volatility_daily', 'Market_Cap']

model = LSTMModel(file_path=file_path, sheet_name=sheet_name, feature_columns=feature_columns)

model.check_exist_data()







































