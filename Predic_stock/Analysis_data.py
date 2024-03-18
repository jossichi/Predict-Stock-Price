from io import BytesIO
import os
from datetime import datetime, timedelta
import re
import docx
from matplotlib.pylab import RandomState
import pandas as pd
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import TimeSeriesSplit
from sklearn.metrics import mean_squared_error
import numpy as np
from datetime import datetime, timedelta
from keras.layers import LSTM, Dense
import matplotlib.pyplot as plt
from keras.models import Sequential
from docx import Document

class Analysis_data:
    def __init__(self, file_path, sheet_name, target_column='Close', feature_columns=None, test_size=0.2, current_date=None, learning_rate=0.001 , n_iterations=100, random_state=None):
        self.file_path = file_path
        self.sheet_name = sheet_name
        self.target_column = target_column
        self.feature_columns = feature_columns
        self.test_size = test_size
        self.current_date = current_date or datetime.now().strftime('%Y-%m-%d')
        self.df = {}  
        self.analysis_results = {}
        self.next_day = (datetime.now() + timedelta(days=1)).strftime('%Y-%m-%d')
        self.model = None
        self.merged_df = None
        self.X = None
        self.y = None
        self.model = None
        self.df_spliting = {}
        self.df_training = []
        self.learning_rate = learning_rate
        self.n_iterations = n_iterations
        self.weights = None
        self.bias = None
        self.random_state = random_state  

    def analyze_data(self):
        if os.path.exists(self.file_path):
            print(f"File '{self.file_path}' tồn tại.")

            # Đọc thông tin về các sheets trong file Excel
            excel_sheets = pd.ExcelFile(self.file_path).sheet_names

            # In ra số lượng sheets và tên của từng sheet
            print(f"Số lượng sheets: {len(excel_sheets)}")
            print(f"Danh sách các sheets: {excel_sheets}")
            for sheet_name in excel_sheets:
                print(f"- {sheet_name}")

                # Đọc dữ liệu từ mỗi sheet
                sheet_df = pd.read_excel(self.file_path, sheet_name)
                # Lấy danh sách các mã cổ phiếu từ cột 'symbols'
                symbols_list = sheet_df['Symbol'].unique()

                # Duyệt qua từng mã cổ phiếu và thực hiện phân tích
                for symbol in symbols_list:
                    print(f"Đang phân tích mã cổ phiếu '{symbol}' trong sheet '{sheet_name}'")
                    
                    # Lọc dữ liệu chỉ chứa thông tin của mã cổ phiếu cụ thể
                    symbol_data = sheet_df[sheet_df['Symbol'] == symbol]

                    # Thực hiện phân tích sử dụng lý thuyết Dow
                    dow_analysis_result, explanation = self.dow_analysis(symbol_data)

                    # Lưu kết quả phân tích vào dictionary
                    self.analysis_results[f"{sheet_name}_{symbol}"] = dow_analysis_result
                    
                    fib_levels,fib_fans,arc_levels,img_stream_retracement,img_stream_fans = self.fibonacci_analysis(symbol_data)
                    
                    # action = self.make_trading_decision(symbol_data, dow_analysis_result)
                    
                    # Xuất kết quả vào file Word
                    self.export_to_word(f"{sheet_name}_{symbol}_analysis_result", dow_analysis_result, explanation,fib_levels,fib_fans,arc_levels,img_stream_retracement,img_stream_fans)
                    
                normalization = False
                if not sheet_df.empty:
                    print(f"\nHandling sheet '{sheet_name}':")

                    missing_values = sheet_df.isnull().sum()
                    nan_values = sheet_df.isna().sum()
                    errors = None

                    if missing_values.any():
                        print(f"\nMissing Values in sheet '{sheet_name}':")
                        print(missing_values)

                        # Xử lý dữ liệu thiếu nếu cần
                        self.handle_missing_values(sheet_df)

                    if normalization:
                        print(f"Data in sheet '{sheet_name}' needs normalization.")
                        self.check_normalization(sheet_name)
                        self.normalize_stock_prices(sheet_name)

                    if nan_values.any():
                        print(f"\nNaN Values in sheet '{sheet_name}':")
                        print(nan_values)

                        # Xử lý NaN values nếu cần
                        self.handle_nan_values(sheet_df)
                        #sau khi xử lý xong các giá trị nan thì lưu vào dataframe
                        self.df[sheet_name] = sheet_df
                        print(f"dataframe xử lý sơ bộ: {self.df}")
                    
                    if errors:
                        print(f"\nErrors in sheet '{sheet_name}':")
                        print(errors)

                        # Xử lý errors nếu cần
                        self.handle_errors(sheet_df)
                        
                    # Lấy danh sách các mã cổ phiếu từ cột 'symbols'
                    symbols_list = sheet_df['Symbol'].unique()

                    # Duyệt qua từng mã cổ phiếu và thực hiện phân tích
                    for symbol in symbols_list:
                    # Update self.df with processed data
                        if nan_values.any():
                            print(f"\nNaN Values in sheet '{sheet_name}':")
                            print(nan_values)

                            # Xử lý NaN values nếu cần
                            self.handle_nan_values(sheet_df)
                            self.df[sheet_name] = sheet_df
                            print(f"dataframe xử lý sơ bộ lần 2: {self.df}")
                    
            print("Tiền xử lý dữ liệu hoàn tất")
            print(f"dataframe cuối cùng: {self.df}")
            self.export_processed_data_to_excel(output_file_path='processed_stock_data.xlsx')

        elif not self.df:
            print("Không có dữ liệu để kiểm tra.")
            return
    
    def dow_analysis(self, data):
        # Đầu tiên, chúng ta cần xác định xu hướng chính của thị trường
        # Điều này có thể được thực hiện bằng cách sử dụng một đường xu hướng tuyến tính hoặc một đường trung bình di động
        trend = self.identify_trend(data)

        # Tiếp theo, chúng ta cần xác định các mức hỗ trợ và kháng cự
        # Điều này có thể được thực hiện bằng cách tìm các điểm cao và thấp trong dữ liệu
        support_level, resistance_level = self.identify_support_resistance(data)

        # Cuối cùng, chúng ta cần xác định xem thị trường có đang trong một giai đoạn tích lũy hay phân phối không
        # Điều này có thể được thực hiện bằng cách phân tích biểu đồ giá và khối lượng giao dịch
        accumulation_distribution = self.identify_accumulation_distribution(data)

        dynamic_support, dynamic_resistance = self.identify_dynamic_support_resistance(data)

        minor_trend = self.identify_minor_trend(data)
        
        # reversal_pattern, continuation_pattern = self.identify_chart_patterns(data)
        
        # confirmation_condition = self.confirm_trend(data, trend)
        
        
        # Kết quả phân tích
        analysis_result = {
            'trend': trend,
            'support_level': support_level,
            'resistance_level': resistance_level,
            'accumulation_distribution': accumulation_distribution,
            'minor trend': minor_trend,
            'dynamic_support': dynamic_support,
            'dynamic_resistance': dynamic_resistance,
            # 'reversal_pattern': reversal_pattern,
            # 'continuation_pattern': continuation_pattern,
        }

        # Gọi hàm explain_analysis_result với kết quả phân tích
        explanation = self.explain_analysis_result(analysis_result)

        return analysis_result, explanation

    def identify_trend(self, data):
        # Sử dụng đường trung bình di động để xác định xu hướng
        # Điều này chỉ là một ví dụ đơn giản, bạn có thể cải thiện nó tùy thuộc vào chiến lược phân tích của bạn
        moving_average = data['Close'].rolling(window=20).mean()

        if data['Close'].iloc[-1] > moving_average.iloc[-1]:
            uptrend_condition = self.confirm_trend(data, 'Uptrend')
            return 'Uptrend' if uptrend_condition else 'No Trend'
        elif data['Close'].iloc[-1] < moving_average.iloc[-1]:
            downtrend_condition = self.confirm_trend(data, 'Downtrend')
            return 'Downtrend' if downtrend_condition else 'No Trend'
        else:
            return 'Sideways'

    def fibonacci_analysis(self,data):
        fib_levels,img_stream_retracement = self.fibonacci_retracement(data)
        fib_fans,img_stream_fans = self.fibonacci_fans(data)
        arc_levels = self.fibonacci_arcs(data)
        return fib_levels,fib_fans,arc_levels,img_stream_retracement,img_stream_fans
    
    def identify_support_resistance(self, data):
        # Tìm giá trị thấp nhất (hỗ trợ) và cao nhất (kháng cự) trong dữ liệu
        support_level = data['Low'].min()
        resistance_level = data['High'].max()

        return support_level, resistance_level

    def identify_accumulation_distribution(self, data):
        # Thực hiện phân tích dựa trên biểu đồ giá và khối lượng giao dịch
        # Điều này chỉ là một ví dụ đơn giản, bạn có thể sử dụng các chỉ báo khác như On-Balance Volume (OBV) để cải thiện
        avg_price = (data['High'] + data['Low']) / 2
        daily_range = data['High'] - data['Low']

        price_change = avg_price.diff()
        volume_change = data['Volume'].diff()

        money_flow = avg_price * data['Volume'] * (price_change / daily_range)
        cumulative_money_flow = money_flow.cumsum()

        if cumulative_money_flow.iloc[-1] > cumulative_money_flow.iloc[0]:
            return 'Accumulation'
        else:
            return 'Distribution'

    def identify_minor_trend(self, data):
        # Sử dụng đường trung bình di động ngắn hạn để xác định xu hướng phụ
        short_term_moving_average = data['Close'].rolling(window=5).mean()

        if data['Close'].iloc[-1] > short_term_moving_average.iloc[-1]:
            return 'Uptrend'
        elif data['Close'].iloc[-1] < short_term_moving_average.iloc[-1]:
            return 'Downtrend'
        else:
            return 'Sideways'

    def identify_dynamic_support_resistance(self, data):
        # Sử dụng Bollinger Bands để xác định mức hỗ trợ và kháng cự
        window = 20  # Độ dài cửa sổ
        std_multiplier = 2  # Hệ số độ lệch chuẩn

        # Tính giá trung bình và độ lệch chuẩn
        data['rolling_mean'] = data['Close'].rolling(window=window).mean()
        data['upper_band'] = data['rolling_mean'] + std_multiplier * data['Close'].rolling(window=window).std()
        data['lower_band'] = data['rolling_mean'] - std_multiplier * data['Close'].rolling(window=window).std()

        # Xác định mức hỗ trợ và kháng cự
        support_level = data['lower_band'].min()
        resistance_level = data['upper_band'].max()

        return support_level, resistance_level

    def identify_chart_patterns(self, data):
        # Xác định các mẫu đảo chiều và tiếp tục xu hướng
        reversal_pattern = self.identify_reversal_pattern(data)
        continuation_pattern = self.identify_continuation_pattern(data)

        return reversal_pattern, continuation_pattern

    def identify_reversal_pattern(self, data):
        # Chức năng này có thể kiểm tra các điểm đảo chiều phổ biến như đầu vai đỉnh, đầu vai đáy, tam giác, ...

        # Ví dụ đơn giản: Kiểm tra xem giá có tạo đỉnh đầu vai không
        if len(data) >= 3:
                # Check for the reversal pattern (e.g., Head and Shoulders)
                if (
                    data['High'].iloc[-3] < data['High'].iloc[-2] > data['High'].iloc[-1] and
                    data['Low'].iloc[-3] < data['Low'].iloc[-2] > data['Low'].iloc[-1]
                ):
                    return 'Head and Shoulders Reversal Pattern'
            
            # If not enough data or no reversal pattern detected, return None
        return None

    def identify_continuation_pattern(self, data):
        # Chức năng này có thể kiểm tra các mẫu tiếp tục xu hướng như cờ, cánh cổ bò, hình chóp nghỉ, ...

        # Ví dụ đơn giản: Kiểm tra xem có mẫu cờ không
        if (
            data['High'].iloc[-3] < data['High'].iloc[-2] > data['High'].iloc[-1] and
            data['Low'].iloc[-3] > data['Low'].iloc[-2] < data['Low'].iloc[-1]
        ):
            return 'Flag Continuation Pattern'
        else:
            return None

    def confirm_trend(self, data, trend_type):
        # Kiểm tra điều kiện xác nhận theo lý thuyết Dow
        confirmation_condition = False

        if trend_type == 'Uptrend':
            # Kiểm tra nến tăng giá với khối lượng tăng
            last_candle = data.iloc[-1]
            previous_candle = data.iloc[-2]

            if last_candle['Close'] > previous_candle['Close'] and last_candle['Volume'] > previous_candle['Volume']:
                confirmation_condition = True

        elif trend_type == 'Downtrend':
            # Kiểm tra nến giảm giá với khối lượng tăng
            last_candle = data.iloc[-1]
            previous_candle = data.iloc[-2]

            if last_candle['Close'] < previous_candle['Close'] and last_candle['Volume'] > previous_candle['Volume']:
                confirmation_condition = True

        return confirmation_condition
    
    def explain_analysis_result(self, analysis_result):
        explanation = "Chi tiết phân tích kết quả:\n"

        # Xu hướng chính
        explanation += f"\nXu hướng chính: {analysis_result['trend']}\n"
        explanation += "\nXu hướng chính mô tả hướng di chuyển chung của thị trường. 'Uptrend' có nghĩa là giá có xu hướng tăng, 'Downtrend' là giá có xu hướng giảm, 'Sideways' là giá dao động trong khoảng nhỏ.\n"

        # Mức hỗ trợ và kháng cự
        explanation += f"\nMức hỗ trợ: {analysis_result['support_level']}\n"
        explanation += f"\nMức hỗ trợ là mức giá thấp nhất mà thị trường thường không giảm thêm. Nếu giá chạm mức hỗ trợ, có thể có sự hỗ trợ tăng giá hoặc ngừng giảm giá.\n"

        explanation += f"\nMức kháng cự: {analysis_result['resistance_level']}\n"
        explanation += f"\nMức kháng cự là mức giá cao nhất mà thị trường thường không vượt qua được. Nếu giá chạm mức kháng cự, có thể có sự kháng cự giảm giá hoặc ngừng tăng giá.\n"

        # Giai đoạn tích lũy hoặc phân phối
        explanation += f"\nGiai đoạn tích lũy hoặc phân phối: {analysis_result['accumulation_distribution']}\n"
        explanation += "\nGiai đoạn tích lũy có thể chỉ ra sự chấp nhận giá và chuẩn bị cho xu hướng tăng giá. Giai đoạn phân phối có thể chỉ ra sự không chấp nhận giá và chuẩn bị cho xu hướng giảm giá.\n"

        # Xu hướng phụ
        explanation += f"\nXu hướng phụ: {analysis_result['minor trend']}\n"
        explanation += "\nXu hướng phụ là xu hướng ngắn hạn trong thời gian ngắn. Nó có thể cung cấp thông tin về sự biến động ngắn hạn của thị trường.\n"

        # Hỗ trợ động và kháng cự động
        explanation += f"\nHỗ trợ động: {analysis_result['dynamic_support']}\n"
        explanation += f"\nGiá trị hỗ trợ động được xác định bằng Bollinger Bands, thường là giá trung bình động trừ một độ lệch chuẩn. Nó có thể cung cấp thông tin về mức giá có thể coi là an toàn.\n"

        explanation += f"\nKháng cự động: {analysis_result['dynamic_resistance']}\n"
        explanation += f"\nGiá trị kháng cự động được xác định bằng Bollinger Bands, thường là giá trung bình động cộng một độ lệch chuẩn. Nó có thể cung cấp thông tin về mức giá có thể gặp khó khăn khi vượt qua.\n"

        # Mẫu đảo chiều và tiếp tục xu hướng
        # explanation += f"\nMẫu đảo chiều: {analysis_result['reversal_pattern']}\n"
        # explanation += f"Mẫu đảo chiều là các biểu hiện trên biểu đồ cho thấy khả năng đảo chiều xu hướng hiện tại. Ví dụ: {analysis_result['reversal_pattern']}."

        # explanation += f"\nMẫu tiếp tục xu hướng: {analysis_result['continuation_pattern']}\n"
        # explanation += f"Mẫu tiếp tục xu hướng là các biểu hiện trên biểu đồ cho thấy khả năng tiếp tục xu hướng hiện tại. Ví dụ: {analysis_result['continuation_pattern']}."

        return explanation

    def export_to_word(self, sheet_name, analysis_result, explanation, fib_levels,fan_level,arc_levels,img_stream_retracement,img_stream_fans):
        # Tạo hoặc mở tệp Word
        doc = Document()

        # Thêm tiêu đề
        doc.add_heading(f"Phân tích kỹ thuật chứng khoán - {sheet_name}", level=1)

        # Thêm thông tin phân tích
        for key, value in analysis_result.items():
            doc.add_paragraph(f"{key}: {value}")

        # Thêm phần giải thích
        doc.add_heading("Giải thích", level=2)
        doc.add_paragraph(explanation)

        # Thêm Fibonacci retracement levels vào giải thích
        doc.add_heading("Fibonacci Retracement Levels", level=3)
        doc.add_paragraph("Các mức Fibonacci retracement:")
        for i, level in enumerate(fib_levels, start=1):
            doc.add_paragraph(f"Level {i}: {level}")
        # Save the plot as an image
        doc.add_picture(img_stream_retracement, width=docx.shared.Inches(5))

        
        doc.add_paragraph("Các mức Fibonacci fans:")
        for i, level in enumerate(fan_level, start=1):
            doc.add_paragraph(f"Level {i}: {level}")
        
        doc.add_picture(img_stream_fans, width=docx.shared.Inches(5))
        
        doc.add_paragraph("Các mức Fibonacci arc:")
        for i, level in enumerate(arc_levels, start=1):
            doc.add_paragraph(f"Level {i}: {level}")

        # doc.add_heading("Trading Decision", level=2)
        # doc.add_paragraph(action)
        # Lưu tệp Word
        output_path = f"report/{sheet_name}_analysis_result.docx"
        doc.save(output_path)
        print(f"Kết quả phân tích đã được lưu vào file Word: {output_path}")

    def check_normalization(self):
        # Kiểm tra xem dữ liệu đã được chuẩn hóa chưa
        if not self.df:
            print("Không có dữ liệu để kiểm tra.")
            return

        # Chọn các đặc trưng cần kiểm tra
        columns_to_check = ['Open', 'High', 'Low', 'Volume', 'RSI', 'MACD_diff', 'BB_upper', 'BB_lower', 'Returns_daily', 'Volatility_daily', 'Market_Cap']

        # Thống kê trung bình và độ lệch chuẩn
        stats_before_normalization = self.df[columns_to_check].describe().loc[['mean', 'std']]
        
        print("\nStatistics Before Normalization:")
        print(stats_before_normalization)

    def normalize_stock_prices(self):
        # Chọn các đặc trưng cần chuẩn hóa
        columns_to_normalize = ['Open', 'High', 'Low', 'Volume', 'RSI', 'MACD_diff', 'BB_upper', 'BB_lower', 'Returns_daily', 'Volatility_daily', 'Market_Cap']

        # Kiểm tra xem có dữ liệu để chuẩn hóa không
        if not self.df:
            print("Không có dữ liệu để chuẩn hóa.")
            return

        # Tạo một bản sao của dữ liệu để tránh ảnh hưởng đến dữ liệu gốc
        normalized_df = self.df.copy()

        # Tiêu chuẩn hóa dữ liệu sử dụng MinMaxScaler
        scaler = MinMaxScaler(feature_range=(-1, 1))
        normalized_df[columns_to_normalize] = scaler.fit_transform(normalized_df[columns_to_normalize])

        # In ra một số dòng đầu của dữ liệu đã được chuẩn hóa
        print("\nNormalized Stock Prices:")
        print(normalized_df.head())

        # Cập nhật dữ liệu trong đối tượng ADALINE
        self.df = normalized_df
    
    def remove_empty_sheets(self):
        empty_sheets = []
        for sheet_name, sheet_df in self.df.items():
            if sheet_df.empty:
                empty_sheets.append(sheet_name)
                print(f"Sheet '{sheet_name}' has no data and will be removed.")

        # Xóa các sheet chỉ có header
        if empty_sheets:
            for empty_sheet in empty_sheets:
                del self.df[empty_sheet]
            print("Empty sheets have been removed.")
    
    def handle_missing_values(self, sheet_df):
        # Kiểm tra và xử lý dữ liệu thiếu
        missing_columns = ['RSI', 'MACD_diff', 'BB_upper', 'BB_lower', 'Returns_daily', 'Volatility_daily', 'Market_Cap']
        missing_data = sheet_df[missing_columns]

        # Hiển thị dữ liệu trước khi xử lý
        print("Missing Data Before Handling:")
        print(missing_data)

        # Thay thế các giá trị thiếu bằng giá trị trung bình của cột
        sheet_df[missing_columns] = sheet_df[missing_columns].fillna(sheet_df[missing_columns].mean())

        # Hiển thị dữ liệu sau khi xử lý
        print("\nMissing Data After Handling:")
        print(sheet_df[missing_columns])

    def handle_nan_values(self,sheet_df):
        sheet_df.ffill(inplace=True)  # Sử dụng forward-fill để điền giá trị NaN từ giá trị trước đó
        sheet_df.bfill(inplace=True)  # Sử dụng backward-fill để điền giá trị NaN từ giá trị sau đó

        # Kiểm tra lại sau khi xử lý
        nan_values_after_handling = sheet_df.isna().sum()

        if nan_values_after_handling.any():
            print(f"\nNaN Values After Handling in sheet: {nan_values_after_handling}")
        else:
            print(f"\nNaN Values have been successfully handled in sheet.")

    def handle_errors(self,sheet_df):
        # Thêm logic xử lý errors nếu cần
        pass
    
    def export_processed_data_to_excel(self, output_file_path='processed_stock_data.xlsx'):   
        with pd.ExcelWriter(output_file_path, engine='xlsxwriter') as writer:
            for sheet_name, sheet_df in self.df.items():
                # Export each sheet's data to a separate sheet in the Excel file
                sheet_df.to_excel(writer, sheet_name=sheet_name, index=False)
                print(f"\nProcessed data for sheet '{sheet_name}' has been exported.")
        
        print(f"\nAll processed data has been exported to '{output_file_path}'.")

    def fibonacci_retracement(self, data):
        # Assuming 'Close' is the column you want to use for the analysis
        close_prices = data['Close']

        # Calculate the highest and lowest price
        high_price = np.max(close_prices)
        low_price = np.min(close_prices)

        # Define Fibonacci retracement levels (commonly used are 38.2%, 50%, and 61.8%)
        fib_levels = [0.382, 0.5, 0.618]

        # Calculate retracement levels
        retracement_levels = [high_price - level * (high_price - low_price) for level in fib_levels]

        # Illustrate Fibonacci retracement levels on a plot
        plt.figure(figsize=(10, 6))
        plt.plot(data['Datetime'], close_prices, label='Close Prices')
        for level, retracement_level in zip(fib_levels, retracement_levels):
            plt.axhline(retracement_level, linestyle='--', color='r', label=f'Fib {level * 100}%')

        plt.title('Fibonacci Retracement Levels')
        plt.xlabel('Datetime')
        plt.ylabel('Close Prices')
        plt.legend()

        img_stream_retracement = BytesIO()
        plt.savefig(img_stream_retracement, format='png')
        img_stream_retracement.seek(0)

        # Close the Matplotlib plot
        plt.close()

        return retracement_levels,img_stream_retracement
   

    def fibonacci_fans(self, data):
        # Assuming 'Close' is the column you want to use for the analysis
        close_prices = data['Close']
        dates = pd.to_datetime(data['Datetime'])  # Convert 'Datetime' column to datetime type

        # Calculate the highest and lowest price
        high_price = np.max(close_prices)
        low_price = np.min(close_prices)

        # Get the dates corresponding to the highest and lowest price
        high_date = dates[close_prices.idxmax()]
        low_date = dates[close_prices.idxmin()]

        # Define Fibonacci retracement levels (commonly used are 23.6%, 38.2%, 50%, 61.8%, and 100%)
        fib_levels = [0.236, 0.382, 0.5, 0.618, 1.0]

        # Calculate fan levels
        fan_levels = [(high_date, high_price - level * (high_price - low_price)) for level in fib_levels]

        # Ensure fan levels are numeric
        fan_levels = [(date, float(price)) for date, price in fan_levels if isinstance(price, (int, float))]

        # Plot the Fibonacci Fans
        plt.figure(figsize=(10, 6))
        plt.plot(dates, close_prices, label='Close Price')
        for level, (date, price) in zip(fib_levels, fan_levels):
            plt.plot([date, dates.iloc[-1]], [high_price, price], label=f'Fan Level {int(level * 100)}%')

        plt.title('Fibonacci Fans Analysis')
        plt.xlabel('Date')
        plt.ylabel('Price')
        plt.legend()

        img_stream_fans = BytesIO()
        plt.savefig(img_stream_fans, format='png')
        img_stream_fans.seek(0)
        plt.close()

        return fan_levels, img_stream_fans

    def fibonacci_arcs(self,data):
        # Assuming 'Close' is the column you want to use for the analysis
        close_prices = data['Close']

        # Calculate the price range (high - low) for the given data
        price_range = np.max(close_prices) - np.min(close_prices)

        fib_levels = [0.382, 0.5, 0.618]

        # Calculate arc levels as percentages
        arc_levels = [np.min(close_prices) + level * price_range for level in fib_levels]

        return arc_levels
    
   
def main():
    # Set file path and other parameters
    file_path = "stock_data.xlsx"
    sheet_name = None
    feature_columns = ['Open', 'High', 'Low', 'Volume', 'RSI', 'MACD_diff', 'BB_upper', 'BB_lower', 'Returns_daily', 'Volatility_daily', 'Market_Cap']

    # Initialize the Analysis_data object
    model = Analysis_data(file_path=file_path, sheet_name=sheet_name, feature_columns=feature_columns, random_state=42)

    # Load and analyze data
    model.analyze_data()

    
if __name__ == "__main__":
    main()
